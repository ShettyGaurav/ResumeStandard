import pdfplumber
from typing_extensions import TypedDict, Annotated, List, Dict, Any
import operator
from dotenv import load_dotenv
import json
from collections import defaultdict
import pprint
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from main import generate_resume_pdf
import os
from markitdown import MarkItDown
import streamlit as st


from langchain.agents import create_agent
from langchain.messages import SystemMessage, HumanMessage
from langgraph.graph import StateGraph, START, END
from langchain_groq import ChatGroq

load_dotenv()

# client = genai.Client()
# model_name = "gemini-2.0-flash"


class State(TypedDict):
    name: str
    education: Annotated[List[str], operator.add]
    experience: str
    place: str
    skills: Dict[str, Any]
    file_path: str
    reference_file_path: str
    content: str
    parse_data: str
    classified_pages: str
    sections: Annotated[List[str], operator.add]


llm = ChatGroq(
    model="llama-3.3-70b-versatile",
    temperature=0,
    max_tokens=None,
    max_retries=2,
)


def get_content(state: State):
    try:
        print("[INFO] Starting file text extraction")
        file_path = state.get("file_path")
        print(f"[INFO] File path received: {file_path}")

        if not file_path or not os.path.exists(file_path):
            print("[ERROR] File path is invalid or file does not exist")
            return {**state, "error": "Invalid file path", "status": "FAILED"}

        text_content = []

        # -------- PDF --------
        if file_path.lower().endswith(".pdf"):
            print("[INFO] Detected PDF file")
            with pdfplumber.open(file_path) as pdf:
                print(f"[INFO] Total pages found: {len(pdf.pages)}")
                for idx, page in enumerate(pdf.pages, start=1):
                    print(f"[INFO] Extracting text from PDF page {idx}")
                    page_text = page.extract_text()
                    if page_text:
                        print(
                            f"[INFO] Text extracted from page {idx} (length: {len(page_text)})"
                        )
                        text_content.append(page_text)
                    else:
                        print(f"[WARN] No text found on page {idx}")

        # -------- DOCX --------
        elif file_path.lower().endswith(".docx"):
            print("[INFO] Detected DOCX file")
            doc = Document(file_path)
            print(f"[INFO] Total paragraphs found: {len(doc.paragraphs)}")
            for idx, para in enumerate(doc.paragraphs, start=1):
                if para.text.strip():
                    print(
                        f"[INFO] Extracting paragraph {idx} (length: {len(para.text)})"
                    )
                    text_content.append(para.text)
                else:
                    print(f"[WARN] Skipping empty paragraph {idx}")

        else:
            print("[ERROR] Unsupported file format")
            return {**state, "error": "Unsupported file format", "status": "FAILED"}

        combined_text = "\n".join(text_content)
        print(combined_text)
        print(f"[INFO] Total extracted content length: {len(combined_text)}")

        state["content"] = combined_text
        print("[INFO] Text extraction completed successfully")

        return state
    except Exception as e:
        return {**state, "error": str(e), "status": "FAILED"}


def get_content_markdown(state: State):
    file_path = state.get("file_path")
    md = MarkItDown(enable_plugins=False)  # Set to True to enable plugins
    # if file_path.endswith(".jpg"):
    #     md = MarkItDown(llm_client=client, llm_model=model_name)
    result = md.convert(file_path)
    print("Markdown content started---------------------------------------")
    print(result.text_content)
    state["content"] = result.text_content
    print("Markdown content ended---------------------------------------")
    urls = []
    markdown_text = result.text_content
    # ---- PDF URLs ----
    if file_path.lower().endswith(".pdf"):
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                if page.annots:
                    for annot in page.annots:
                        uri = annot.get("uri")
                        if uri:
                            urls.append(uri)

    elif file_path.lower().endswith(".docx"):
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        doc = Document(file_path)
        for rel in doc.part.rels.values():
            if rel.reltype == RT.HYPERLINK:
                urls.append(rel.target_ref)

    if urls:
        markdown_text += "\n\n---\n**Links found in document:**\n"
        for url in sorted(set(urls)):
            markdown_text += f"- {url}\n"
    print(urls)
    state["content"] = markdown_text
    return state


def get_content_strutured(state: State):
    system_prompt = """You are an expert Resume Information Extraction and Normalization Agent.

    Your task is to extract structured resume information from raw resume text and return it in EXACTLY the JSON format defined below.

    STRICT OUTPUT FORMAT (do not change keys outside Skillset, do not add new top-level keys):

    {
    "resume": {
    "name": "string",
    "summary": [
    "string"
    ],
    "sections": {
    "Career Summary": [
    "string"
    ],
    "Professional History": [
    {
    "title": "string",
    "company": "string",
    "timespan": "string",
    "points": [
    "string"
    ]
    }
    ],
    "Skillset": {
    },
    "Project Showcase": [
    {
    "title": "string",
    "technologies": ["string"],
    "points": [
    "string"
    ]
    }
    ],
    "Education": [
    "string"
    ]
    }
    }
    }

    ──────────────── EXTRACTION RULES ────────────────

    NAME

    Extract the candidate’s full name from the top of the resume.

    If not found, return an empty string.

    SUMMARY

    Extract 2–4 concise sentences describing the professional profile.

    Do NOT include achievements, metrics, years, company names, or education.

    Normalize wording while preserving meaning.

    ──────────────── CONTACT DETAILS (MANDATORY) ────────────────

    You must always return a contact object with all keys present.

    If a value is missing, return "None" (string).

    Phone

    Extract phone number if present.

    Normalize spacing.

    If not found → "None".

    Email

    Extract email if present.

    Normalize to lowercase.

    If not found → "None".

    LinkedIn
    The resume may contain:

    Full URL
    https://linkedin.com/in/username

    Partial URL
    linkedin.com/in/username

    Username only
    username

    Text like
    LinkedIn: username

    Normalization rule:

    If only a username is found, convert it to:
    https://linkedin.com/in/{{username}}

    If nothing related to LinkedIn exists → "None".

    GitHub
    The resume may contain:

    Full URL
    https://github.com/username

    Username only
    username

    Text like
    GitHub: username

    Normalization rule:

    If only a username is found, convert it to:
    https://github.com/{{username}}

    If nothing related to GitHub exists → "None".

    Location

    Extract city, state, or country if explicitly present.

    Do NOT infer or guess.

    If not found → "None".

    ❗ Do NOT hallucinate any contact information.

    CAREER SUMMARY

    Convert experience descriptions into bullet points.

    Each bullet must be a single concise sentence.

    Do NOT include dates, company names, or education.

    EDUCATION

    Extract only formal academic degrees.

    Include degree, institution, and location if present.

    Do NOT include:

    Courses

    Certifications

    Training programs

    Do NOT include years.

    If no valid academic education exists, return an empty array.

    PROFESSIONAL HISTORY

    Extract work experience with job titles, companies, and time spans.

    Each entry must include:

    title (job position)

    company (organization name)

    timespan (employment duration, e.g., "Jan 2023 - Present")

    points (array of achievements/responsibilities)

    Include specific dates when available.

    If no work experience exists, return an empty array.

    PROJECT SHOWCASE

    Each project must include:

    title

    technologies (array)

    points (array of complete sentences)

    Projects must represent real implementations.

    If no valid projects exist, return an empty array.

    ──────────────── SKILLSET (HIERARCHICAL & FLEXIBLE) ────────────────

    Skillset must be a nested object with logical parent domains and sub-categories.

    Parent Categories (create only if relevant)

    Examples (not exhaustive):

    UI

    Backend

    AI / ML

    Data

    Cloud

    DevOps

    Tools

    Do NOT create a category if the resume does not support it.

    UI (if applicable)

    Allowed sub-categories:

    Frontend Frameworks

    Animation Libraries

    CSS Libraries

    UI Component Libraries

    Examples:

    Frontend Frameworks: React, Angular, Next.js

    Animation Libraries: Anime.js, Framer Motion

    CSS Libraries: Tailwind CSS, Bootstrap

    UI Component Libraries: ShadCN UI, Material UI

    Backend (if applicable)

    Allowed sub-categories:

    Server Runtime

    Backend Frameworks / Libraries

    Databases

    Caching Systems

    Message Queues

    Rules:

    Databases must be grouped by type inside Databases:

    SQL

    NoSQL

    Vector

    Example:

    Server Runtime: Node.js, Go

    Backend Frameworks / Libraries: Express, Gin, Echo

    Databases:

    SQL: PostgreSQL, CockroachDB

    NoSQL: MongoDB, Cassandra

    Vector: PgVector, Pinecone

    Caching Systems: Redis

    AI / ML (if applicable)

    Allowed sub-categories:

    ML Libraries

    Models / LLMs Used

    MLOps Tools

    Examples:

    ML Libraries: PyTorch, TensorFlow

    Models / LLMs Used: Gemini, Groq, Grok

    MLOps Tools: MLflow, DVC

    Data (if applicable)

    Allowed sub-categories:

    Data Science Libraries

    Visualization Tools

    Examples:

    Data Science Libraries: Pandas, NumPy, Polars

    Visualization Tools: Matplotlib, Seaborn

    Cloud (if applicable)

    Allowed sub-categories:

    Cloud Platforms

    Cloud Services

    Examples:

    Cloud Platforms: AWS, Azure

    Cloud Services: ECS, Lambda, EKS

    DevOps (if applicable)

    Allowed sub-categories:

    CI / CD

    Containerization & Orchestration

    Monitoring & Logging

    Infrastructure Tools

    Examples:

    CI / CD: Jenkins, GitHub Actions

    Containerization & Orchestration: Docker, Kubernetes, Helm

    Monitoring & Logging: Grafana, Prometheus

    Infrastructure Tools: Ansible, Terraform

    Skillset Rules (MANDATORY)

    Do NOT invent skills.
    If in subcategories there are no values Please remote the category 

    Do NOT duplicate skills across categories.

    Normalize names (e.g., PostgresSQL → PostgreSQL).

    Each sub-category value must be an array of strings.

    Nested grouping (e.g., SQL / NoSQL) must be objects with arrays.

    Do NOT output empty parent categories or empty sub-categories.

    Skillset is the only section allowed to have nested objects.

    ──────────────── OUTPUT RULES (MANDATORY) ────────────────

    Output ONLY valid JSON

    No markdown

    No explanations

    No comments

    No trailing commas

    Do not output anything other than the JSON object"""
    if state["content"]:
        print("Its There")
    human_prompt = f"""Extract structured resume information from the following raw resume text:

    {state["content"]}

    """

    try:

        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=human_prompt),
        ]

        response = llm.invoke(messages)
        jsoncontent = response.content

        if "json" or "```" in jsoncontent.lower():
            jsoncontent = jsoncontent.strip("```json").strip("```")

        jsoncontent = jsoncontent.strip()
        print("=================== Raw Response from LLM =======================")
        print(response.content)
        print("Json content Conversion Start")
        jsonResponse = json.loads(jsoncontent)
        print("Json content Conversion End")
        state["parse_data"] = jsonResponse
        pprint.pprint(state["parse_data"])
        return state
    except Exception as e:
        return f"Error extracting content: {str(e)}"


def generate_PDF(state: State):
    print("Called ")
    try:
        result = generate_resume_pdf(state["parse_data"], show_contact=True)
        print(result)
        return state
    except Exception as e:
        return f"Error extracting content: {str(e)}"


workflow = StateGraph(State)

workflow.add_node("generate_pdf", generate_PDF)
workflow.add_node("get_content", get_content)
workflow.add_node("get_content_markdown", get_content_markdown)
workflow.add_node("get_content_structured", get_content_strutured)
# workflow.add_node("get_experience", get_experience)
# workflow.add_node("get_sections", get_sections)

workflow.add_edge(START, "get_content_markdown")
workflow.add_edge("get_content_markdown", "get_content_structured")
workflow.add_edge("get_content_structured", "generate_pdf")
workflow.add_edge("generate_pdf", END)


# workflow.add_edge("get_content_markdown",END)

graph = workflow.compile()


# response = graph.invoke(
#     {
#         "file_path": "Resume/Karthik_Resume.pdf",
#     }
# )

# response = graph.invoke(
#     {
#         "file_path": "ResumeFolder/Gauravv.pdf",
#     }
# )


INPUT_DIR = "ResumeFolder"


def get_response(file_path) -> str:
    try:
        response = graph.invoke(
            {
                "file_path": file_path,
            }
        )
        print("TRue")
    except Exception as e:
        return f"Error extracting content: {str(e)}"
    return f"the file named {file_path} is created"
